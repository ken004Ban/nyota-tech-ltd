from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

BLUE = RGBColor(0, 82, 155)
DARK_BLUE = RGBColor(0, 50, 110)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = RGBColor(230, 240, 255)

doc = Document()

# -- Page setup --
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BLACK

# ── Helper: horizontal rule ──
def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00529B')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '00529B')
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table_row(table, cells_data, is_header=False):
    row = table.add_row()
    for i, (text, bold, color, align) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.bold = bold
        run.font.color.rgb = color
        set_cell_border(cell)
        if is_header:
            set_cell_shading(cell, '00529B')

# ── COVER PAGE ──
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NYOTA TECH')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Software & Systems')
run.font.size = Pt(14)
run.font.color.rgb = BLUE
run.italic = True

add_hr(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Proposal & Quotation')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WhatsApp Business API Integration')
run.font.size = Pt(16)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for: [Client Name]')
run.font.size = Pt(12)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Date: {datetime.date.today().strftime("%d %B %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Reference: NT-Q-2026-001')
run.font.size = Pt(11)
run.font.color.rgb = BLACK

for _ in range(4):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nyota Tech Limited')
run.font.size = Pt(11)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('nyotatech.info@gmail.com')
run.font.size = Pt(10)
run.font.color.rgb = BLUE

# ── PAGE BREAK ──
doc.add_page_break()

# ── TABLE OF CONTENTS (compact) ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Table of Contents')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'Technical Solution Overview'),
    ('3.', 'Scope of Work'),
    ('4.', 'Pricing & Fees'),
    ('5.', 'Payment Terms'),
    ('6.', 'Timeline'),
    ('7.', 'Client Responsibilities'),
    ('8.', 'Terms & Conditions'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK

doc.add_paragraph('')

# ── 1. EXECUTIVE SUMMARY ──
p = doc.add_paragraph()
run = p.add_run('1.  Executive Summary')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

p = doc.add_paragraph()
run = p.add_run(
    'This proposal outlines an integration of Meta\'s official WhatsApp Cloud API '
    'into [Client]\'s website, enabling the site to send WhatsApp messages to three '
    'designated mobile numbers programmatically — triggered by user form input.'
)
run.font.size = Pt(11)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
run = p.add_run(
    'The solution uses the WhatsApp Cloud API hosted by Meta (no on-premise infrastructure). '
    'A lightweight backend receives website form submissions and dispatches messages to each '
    'of the three recipients via approved message templates. All communications comply with '
    'WhatsApp\'s platform policies, including business verification, template approvals, '
    'and user opt-in consent requirements.'
)
run.font.size = Pt(11)
run.font.color.rgb = BLACK

p = doc.add_paragraph()
run = p.add_run('Target recipient numbers:')
run.bold = True
run.font.size = Pt(11)

for num in ['+260 962 026 492', '+260 771 204 563', '+260 770 097 348']:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(num)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.bold = True

doc.add_paragraph('')

# ── 2. TECHNICAL SOLUTION OVERVIEW ──
p = doc.add_paragraph()
run = p.add_run('2.  Technical Solution Overview')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

p = doc.add_paragraph()
run = p.add_run('Architecture')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

steps = [
    ('Website Form', 'A user fills in a web form on [Client]\'s site and submits it.'),
    ('Backend Receiver', 'The form data is sent to a lightweight backend endpoint (Node.js / Python).'),
    ('WhatsApp Cloud API', 'The backend calls Meta\'s Graph API (POST /messages) three times — once per recipient.'),
    ('Message Delivery', 'Each recipient receives a WhatsApp message via a pre-approved template with personalised content.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}  —  ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run = p.add_run(desc)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

p = doc.add_paragraph()
run = p.add_run('Key Technical Details')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

details = [
    'API Endpoint: POST https://graph.facebook.com/v25.0/{Phone-Number-ID}/messages',
    'Authentication: Permanent system-user token (no 24-hour expiry)',
    'Message types supported: Text, template, media, interactive',
    'Template categories: Utility (updates) and Marketing (promotions)',
    'Meta infrastructure: Fully hosted — no servers, Docker, or on-premise maintenance',
]
for d in details:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(d)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

p = doc.add_paragraph()
run = p.add_run('Platform Policy Compliance')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

compliance = [
    'Business verification via Meta Business Manager (company documentation required)',
    'Pre-approved message templates for all business-initiated outbound messages',
    'Explicit opt-in consent from each recipient before sending (auditable record)',
    '24-hour customer-service window for free-form follow-up replies',
]
for c in compliance:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(c)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

doc.add_paragraph('')

# ── 3. SCOPE OF WORK ──
p = doc.add_paragraph()
run = p.add_run('3.  Scope of Work')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

p = doc.add_paragraph()
run = p.add_run('Phase 1 — Setup (one-time)')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

setup_items = [
    'Meta Business Manager account setup and configuration',
    'Business verification submission and follow-through with Meta',
    'WhatsApp Business Account (WABA) creation',
    'Dedicated phone number registration and display name approval',
    'Message template creation (1 Utility + 1 Marketing template) + Meta approval',
    'User opt-in consent mechanism (web checkbox + record-keeping)',
    'Backend integration — web form capture → API dispatch to 3 numbers',
    'End-to-end testing with all three recipient numbers',
    'Handover documentation and runbook',
]
for item in setup_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

p = doc.add_paragraph()
run = p.add_run('Phase 2 — Maintenance (monthly recurring)')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

maint_items = [
    'Serverless / cloud hosting and uptime monitoring',
    'WhatsApp template management and renewal as needed',
    'Token and credential rotation (security best practice)',
    'Ongoing technical support (48-hour first-response target)',
    'Monthly delivery report',
]
for item in maint_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

doc.add_paragraph('')

# ── 4. PRICING & FEES ──
p = doc.add_paragraph()
run = p.add_run('4.  Pricing & Fees')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

p = doc.add_paragraph()
run = p.add_run('One-Time Setup Fee')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

setup_table = doc.add_table(rows=1, cols=3)
setup_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr = setup_table.rows[0]
for i, (txt, bold, color, align) in enumerate([
    ('Item', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Detail', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Amount (USD)', True, WHITE, WD_ALIGN_PARAGRAPH.RIGHT),
]):
    cell = hdr.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(txt)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    set_cell_border(cell)
    set_cell_shading(cell, '00529B')

setup_rows = [
    ('Base setup per number', '3 numbers × $100 each', '$300.00'),
    ('Project margin', 'Project management, comms, buffer', '$75.00'),
    ('', '', ''),
]
for item, detail, amount in setup_rows:
    cells_data = [
        (item, True if item else False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
        (detail, False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
        (amount, True if amount else False, BLACK, WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    add_table_row(setup_table, cells_data)

# Total row
total_row = setup_table.add_row()
for i, (txt, bold, color, align) in enumerate([
    ('Total One-Time Setup', True, DARK_BLUE, WD_ALIGN_PARAGRAPH.LEFT),
    ('', False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
    ('$375.00', True, DARK_BLUE, WD_ALIGN_PARAGRAPH.RIGHT),
]):
    cell = total_row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(txt)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    set_cell_border(cell)
    set_cell_shading(cell, 'E6F0FF')

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Monthly Maintenance Fee')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

maint_table = doc.add_table(rows=1, cols=3)
maint_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = maint_table.rows[0]
for i, (txt, bold, color, align) in enumerate([
    ('Item', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Detail', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Amount (USD)', True, WHITE, WD_ALIGN_PARAGRAPH.RIGHT),
]):
    cell = hdr.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(txt)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    set_cell_border(cell)
    set_cell_shading(cell, '00529B')

maint_rows = [
    ('Hosting & infrastructure', 'Serverless / cloud hosting', '$20.00'),
    ('Monitoring & support', '48h first-response, template mgmt', '$10.00'),
]
for item, detail, amount in maint_rows:
    cells_data = [
        (item, True, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
        (detail, False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
        (amount, True, BLACK, WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    add_table_row(maint_table, cells_data)

total_row = maint_table.add_row()
for i, (txt, bold, color, align) in enumerate([
    ('Total Monthly', True, DARK_BLUE, WD_ALIGN_PARAGRAPH.LEFT),
    ('', False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
    ('$30.00', True, DARK_BLUE, WD_ALIGN_PARAGRAPH.RIGHT),
]):
    cell = total_row.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(txt)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    set_cell_border(cell)
    set_cell_shading(cell, 'E6F0FF')

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Meta Channel Costs (pass-through, variable)')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DARK_BLUE

meta_table = doc.add_table(rows=1, cols=3)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = meta_table.rows[0]
for i, (txt, bold, color, align) in enumerate([
    ('Category', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Rate per message (+260)', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Billing', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
]):
    cell = hdr.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(txt)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    set_cell_border(cell)
    set_cell_shading(cell, '00529B')

meta_rows = [
    ('Utility (updates, alerts)', '~$0.0040', 'Billed at cost + 5% admin'),
    ('Marketing (promotions)', '~$0.0225', 'Billed at cost + 5% admin'),
    ('Authentication (OTP)', '~$0.0040', 'Billed at cost + 5% admin'),
    ('Service (inbound replies)', 'Free', 'Included'),
]
for item, rate, billing in meta_rows:
    cells_data = [
        (item, False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
        (rate, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.LEFT),
        (billing, False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
    ]
    add_table_row(meta_table, cells_data)

doc.add_paragraph('')

# ── 5. PAYMENT TERMS ──
p = doc.add_paragraph()
run = p.add_run('5.  Payment Terms')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

terms = [
    '50% of the one-time setup fee due upon signed agreement.',
    '50% balance due upon completion, testing, and client acceptance.',
    'Monthly maintenance invoiced at the start of each calendar month.',
    'Meta channel costs billed monthly at cost plus a 5% administrative fee.',
    'Payment is due within 14 days of invoice date.',
    'All prices quoted in United States Dollars (USD).',
]
for t in terms:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

doc.add_paragraph('')

# ── 6. TIMELINE ──
p = doc.add_paragraph()
run = p.add_run('6.  Timeline')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

p = doc.add_paragraph()
run = p.add_run('Estimated duration: 2–3 weeks from signed agreement.')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = DARK_BLUE

timeline_table = doc.add_table(rows=1, cols=3)
timeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = timeline_table.rows[0]
for i, (txt, bold, color, align) in enumerate([
    ('Phase', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Activity', True, WHITE, WD_ALIGN_PARAGRAPH.LEFT),
    ('Days', True, WHITE, WD_ALIGN_PARAGRAPH.RIGHT),
]):
    cell = hdr.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(txt)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    set_cell_border(cell)
    set_cell_shading(cell, '00529B')

timeline_rows = [
    ('1', 'Meta Business setup & verification submission', '1–3'),
    ('2', 'Meta verification processing (outside our control)', '3–14'),
    ('3', 'Template creation & approval (parallel)', '1–3'),
    ('4', 'Backend integration & API wiring', '3–7'),
    ('5', 'End-to-end testing with 3 numbers', '3–5'),
    ('6', 'Go-live & handover', '1–2'),
]
for phase, activity, days in timeline_rows:
    cells_data = [
        (phase, True, DARK_BLUE, WD_ALIGN_PARAGRAPH.CENTER),
        (activity, False, BLACK, WD_ALIGN_PARAGRAPH.LEFT),
        (days, True, BLACK, WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    add_table_row(timeline_table, cells_data)

doc.add_paragraph('')

# ── 7. CLIENT RESPONSIBILITIES ──
p = doc.add_paragraph()
run = p.add_run('7.  Client Responsibilities')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

client_items = [
    'Provide company registration documents for Meta business verification.',
    'Provide a dedicated phone number (not registered on personal WhatsApp).',
    'Ensure opt-in consent is obtained from all three recipient numbers.',
    'Review and approve message template content for Meta submission.',
    'Provide website backend access or coordinate with the existing web developer.',
    'Designate a single point of contact for the duration of the engagement.',
]
for item in client_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

doc.add_paragraph('')

# ── 8. TERMS & CONDITIONS ──
p = doc.add_paragraph()
run = p.add_run('8.  Terms & Conditions')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE

add_hr(doc)

tc_items = [
    'This quotation is valid for 30 days from the date of issue.',
    'Meta\'s business verification timeline is outside the control of Nyota Tech and may affect the overall delivery schedule.',
    'All source code and configuration developed specifically for this project becomes the intellectual property of the client upon full payment.',
    'Nyota Tech is not liable for WhatsApp platform policy changes, account suspensions, or service interruptions caused by Meta.',
    'Any additional feature requests beyond the scope defined in Section 3 will be quoted separately.',
    'Either party may terminate the monthly maintenance agreement with 30 days written notice.',
]
for item in tc_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

doc.add_paragraph('')
add_hr(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nyota Tech Limited  |  nyotatech.info@gmail.com')
run.font.size = Pt(10)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Calm, dependable software.')
run.font.size = Pt(10)
run.font.color.rgb = DARK_BLUE
run.italic = True

# ── SAVE ──
output_path = 'Z:\\nyota-tech-ltd\\docs\\WhatsApp-API-Proposal-Nyota-Tech.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
